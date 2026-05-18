"""Build a Word document from docs/informe-desarrollo-ddd.md plus architecture diagrams."""
from __future__ import annotations

import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.graphics import renderPM
from svglib.svglib import svg2rlg

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "informe-desarrollo-ddd.md"
ARCH_DIR = ROOT / "docs" / "architecture"
OUT_PATH = ROOT / "docs" / "informe-desarrollo-ddd.docx"

IMAGE_SECTIONS = [
    ("c4", "Diagramas C4", [
        ("weatherflow-context.svg", "C4 - Diagrama de Contexto"),
        ("weatherflow-container.svg", "C4 - Diagrama de Contenedores"),
        ("weatherflow-component.svg", "C4 - Diagrama de Componentes"),
    ]),
    ("uml", "Diagramas UML", [
        ("weatherflow-domain-model.svg", "UML - Modelo de Dominio"),
        ("weatherflow-ports-adapters.svg", "UML - Puertos y Adaptadores"),
    ]),
    ("sequences", "Diagramas de Secuencia", [
        ("register-user-sequence.svg", "Secuencia - Registro de Usuario"),
        ("login-user-sequence.svg", "Secuencia - Login de Usuario"),
        ("query-measurements-sequence.svg", "Secuencia - Consulta de Mediciones"),
        ("record-measurement-alert-sequence.svg", "Secuencia - Registro de Medicion con Alerta"),
        ("manage-notification-preferences-sequence.svg", "Secuencia - Gestion de Preferencias de Notificacion"),
    ]),
]


def svg_to_png(svg_path: Path, png_path: Path, scale: float = 1.5) -> None:
    drawing = svg2rlg(str(svg_path))
    if drawing is None:
        raise RuntimeError(f"No se pudo cargar {svg_path}")
    drawing.scale(scale, scale)
    drawing.width *= scale
    drawing.height *= scale
    renderPM.drawToFile(drawing, str(png_path), fmt="PNG")


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            p.add_run(part)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            p.add_run(part)


def add_image(doc: Document, png_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


def parse_markdown(md_text: str):
    blocks = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("- "):
            item = line[2:].strip()
            if item.endswith(";") or item.endswith(":") or item.endswith("."):
                pass
            blocks.append(("bullet", item.rstrip(";")))
            i += 1
        else:
            blocks.append(("p", line.strip()))
            i += 1
    return blocks


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        for kind, content in blocks:
            if kind == "h1":
                add_heading(doc, content, level=0)
            elif kind == "h2":
                add_heading(doc, content, level=1)
            elif kind == "h3":
                add_heading(doc, content, level=2)
            elif kind == "bullet":
                add_bullet(doc, content)
            else:
                add_paragraph(doc, content)

        doc.add_page_break()
        add_heading(doc, "Anexo: Diagramas de Arquitectura", level=1)
        add_paragraph(
            doc,
            "Este anexo contiene los diagramas C4, UML y de secuencia que acompanan el informe de desarrollo. Las imagenes se generaron a partir de los archivos fuente en docs/architecture.",
        )

        for folder, title, images in IMAGE_SECTIONS:
            add_heading(doc, title, level=2)
            for filename, caption in images:
                svg_path = ARCH_DIR / folder / filename
                if not svg_path.exists():
                    add_paragraph(doc, f"[Imagen no encontrada: {svg_path.name}]")
                    continue
                png_path = tmp_path / (svg_path.stem + ".png")
                try:
                    svg_to_png(svg_path, png_path)
                    add_image(doc, png_path, caption)
                except Exception as exc:
                    add_paragraph(doc, f"[Error al convertir {svg_path.name}: {exc}]")

        doc.save(str(OUT_PATH))
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
